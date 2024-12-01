from flask import Flask, request, send_file, render_template
from werkzeug.utils import secure_filename
from PIL import Image
from docx import Document
import os

UPLOAD_FOLDER = './uploads'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

@app.route('/')
def home():
    return render_template('docs.html')

@app.route('/convert-docx-to-image', methods=['POST'])
def convert_docx_to_image():
    try:
        file = request.files['docxFile']
        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)

        doc = Document(filepath)
        text = ''.join(paragraph.text for paragraph in doc.paragraphs)
        binary_data = ''.join(format(ord(char), '08b') for char in text)

        width = 500
        height = (len(binary_data) + width - 1) // width
        img = Image.new('1', (width, height), color=1)

        for i, bit in enumerate(binary_data):
            x = i % width
            y = i // width
            img.putpixel((x, y), 0 if bit == '1' else 1)

        output_path = os.path.join(UPLOAD_FOLDER, 'binary_image.png')
        img.save(output_path, format="PNG")
        return send_file(output_path, mimetype='image/png', as_attachment=True, download_name='binary_image.png')

    except Exception as e:
        return f"Error: {str(e)}", 500

@app.route('/convert-image-to-docx', methods=['POST'])
def convert_image_to_docx():
    try:
        file = request.files['binaryImage']
        filename = secure_filename(file.filename)
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)

        img = Image.open(filepath).convert('1')
        binary_data = ''
        width, height = img.size

        for y in range(height):
            for x in range(width):
                binary_data += '1' if img.getpixel((x, y)) == 0 else '0'

        text = ''.join(chr(int(binary_data[i:i+8], 2)) for i in range(0, len(binary_data), 8))
        doc = Document()
        doc.add_paragraph(text)

        output_path = os.path.join(UPLOAD_FOLDER, 'reconstructed.docx')
        doc.save(output_path)
        return send_file(output_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                         as_attachment=True, download_name='reconstructed.docx')

    except Exception as e:
        return f"Error: {str(e)}", 500

if __name__ == '__main__':
    app.run(debug=True)
